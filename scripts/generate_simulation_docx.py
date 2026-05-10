"""Generate a Word DOCX report with all key results from notebook 04_simulation."""

import sys
from pathlib import Path
import json
import warnings
from datetime import datetime
import tempfile
import os

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "src"))

import numpy as np
import pandas as pd
from sklearn.cluster import KMeans
from sklearn.metrics import adjusted_rand_score
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from features import FINAL_FEATURES
from model_utils import build_cluster_profile, compute_clustering_metrics

warnings.filterwarnings("ignore")

PROCESSED_DIR = PROJECT_ROOT / "data" / "processed"
MODELS_DIR = PROJECT_ROOT / "models"
OUTPUT_DOCX = PROJECT_ROOT / "reports" / "04_simulation_results_clean.docx"
OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

# ── Colors for charts ────────────────────────────────────────────────────────
C_BG = "#ffffff"
C_CARD = "#f8f9fa"
C_ACCENT = "#0f3460"
C_BLUE = "#3498db"
C_GREEN = "#2ecc71"
C_ORANGE = "#f39c12"
C_RED = "#e74c3c"
C_TEXT = "#2c3e50"
C_MUTED = "#95a5a6"

def styled_fig(figsize=(8, 5)):
    fig, ax = plt.subplots(figsize=figsize, facecolor=C_BG)
    ax.set_facecolor(C_BG)
    for spine in ax.spines.values():
        spine.set_color(C_MUTED)
    ax.tick_params(colors=C_TEXT)
    ax.xaxis.label.set_color(C_TEXT)
    ax.yaxis.label.set_color(C_TEXT)
    ax.title.set_color(C_TEXT)
    return fig, ax

def add_table(doc, title, df):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    # Data rows
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    doc.add_paragraph()


def main():
    print("=" * 60)
    print("  GÉNÉRATION DU RAPPORT DOCX — Notebook 04 Simulation")
    print("=" * 60)

    # ── 1. Load data & model ─────────────────────────────────────────────
    df_labeled = pd.read_parquet(PROCESSED_DIR / "customer_features_labeled.parquet")
    # Forcer le modèle de production : UMAP + KMeans k=4
    best_model_path = MODELS_DIR / "best_clustering_kmeans_k4.pkl"
    assert best_model_path.exists(), f"Modèle introuvable : {best_model_path}"
    meta_path = best_model_path.with_suffix(".json")
    with open(meta_path) as f:
        meta = json.load(f)
    BEST_K = 4
    print(f"Modèle : {best_model_path.name} | k={BEST_K}")

    # ── 2. Temporal windows ──────────────────────────────────────────────
    date_col = "last_purchase_date"
    df_labeled[date_col] = pd.to_datetime(df_labeled[date_col])
    df_sim = df_labeled.copy()

    WINDOWS = {
        "T1": pd.Timestamp("2017-06-30"),
        "T2": pd.Timestamp("2017-12-31"),
        "T3": pd.Timestamp("2018-06-30"),
        "T4": df_sim[date_col].max(),
    }
    windows_data = {}
    window_sizes = {}
    for name, cutoff in WINDOWS.items():
        mask = df_sim[date_col] <= cutoff
        windows_data[name] = df_sim[mask].copy()
        window_sizes[name] = int(mask.sum())

    # ── 3. Re-clustering per window ──────────────────────────────────────
    window_labels = {}
    window_customer_ids = {}
    window_metrics = {}

    for name, df_w in windows_data.items():
        X_w = df_w[FINAL_FEATURES].values.astype(np.float64)
        if len(X_w) < BEST_K * 10:
            continue
        km = KMeans(n_clusters=BEST_K, init="k-means++", n_init=20, random_state=42)
        labels = km.fit_predict(X_w)
        window_labels[name] = labels
        window_customer_ids[name] = list(df_w["customer_unique_id"])
        m = compute_clustering_metrics(X_w, labels)
        window_metrics[name] = m

    # ── 4. ARI ───────────────────────────────────────────────────────────
    pairs = [("T1", "T2"), ("T2", "T3"), ("T3", "T4")]
    ari_scores = {}
    ari_common_counts = {}
    for t_prev, t_next in pairs:
        if t_prev not in window_labels or t_next not in window_labels:
            continue
        ids_prev = set(window_customer_ids[t_prev])
        ids_next = set(window_customer_ids[t_next])
        common = sorted(ids_prev & ids_next)
        if len(common) < 50:
            continue
        prev_map = dict(zip(window_customer_ids[t_prev], window_labels[t_prev]))
        next_map = dict(zip(window_customer_ids[t_next], window_labels[t_next]))
        y_prev = np.array([prev_map[c] for c in common])
        y_next = np.array([next_map[c] for c in common])
        ari = adjusted_rand_score(y_prev, y_next)
        ari_scores[f"{t_prev}→{t_next}"] = round(ari, 4)
        ari_common_counts[f"{t_prev}→{t_next}"] = len(common)

    # ── 5. Silhouette by window ──────────────────────────────────────────
    silhouette_by_window = {}
    for name in windows_data:
        if name in window_metrics:
            silhouette_by_window[name] = round(window_metrics[name]["silhouette"], 4)

    # ── 6. Cluster profiles T3 vs T4 ────────────────────────────────────
    RAW_PROFILE_FEATURES = [
        "Monetary", "Frequency", "Frequency_flag",
        "avg_delivery_delay", "avg_review_score", "avg_freight_ratio",
    ]
    raw_cols = [c for c in RAW_PROFILE_FEATURES if c in windows_data["T4"].columns]
    profile_by_window = {}
    for name, df_w in windows_data.items():
        if name not in window_labels or not raw_cols:
            continue
        prof = build_cluster_profile(df_w, window_labels[name], raw_cols)
        profile_by_window[name] = prof

    delta_df = None
    if "T3" in profile_by_window and "T4" in profile_by_window:
        p3 = profile_by_window["T3"].set_index("cluster")[["CLV_proxy"]].rename(
            columns={"CLV_proxy": "CLV_T3"})
        p4 = profile_by_window["T4"].set_index("cluster")[["CLV_proxy"]].rename(
            columns={"CLV_proxy": "CLV_T4"})
        delta_df = p3.join(p4, how="inner")
        delta_df["delta_pct"] = (
            (delta_df["CLV_T4"] - delta_df["CLV_T3"])
            / delta_df["CLV_T3"].abs() * 100
        ).round(2)

    # ── 7. Bootstrap ─────────────────────────────────────────────────────
    N_BOOTSTRAP = 30
    bootstrap_sil = []
    rng = np.random.default_rng(42)
    df_t4 = windows_data.get("T4", df_sim)
    X_t4 = df_t4[FINAL_FEATURES].values.astype(np.float64)
    n_t4 = len(X_t4)

    for i in range(N_BOOTSTRAP):
        idx = rng.choice(n_t4, size=n_t4, replace=True)
        X_boot = X_t4[idx]
        km_boot = KMeans(n_clusters=BEST_K, n_init=5, random_state=i)
        labels_boot = km_boot.fit_predict(X_boot)
        m = compute_clustering_metrics(X_boot, labels_boot)
        if not np.isnan(m["silhouette"]):
            bootstrap_sil.append(m["silhouette"])

    boot_mean = np.mean(bootstrap_sil)
    boot_std = np.std(bootstrap_sil)
    boot_cv = boot_std / boot_mean if boot_mean > 0 else np.nan
    ci95_lo = np.percentile(bootstrap_sil, 2.5)
    ci95_hi = np.percentile(bootstrap_sil, 97.5)

    # ── 8. Signals synthesis ─────────────────────────────────────────────
    signals = []
    ari_t3_t4 = ari_scores.get("T3→T4")
    if ari_t3_t4 is not None:
        if ari_t3_t4 > 0.7:
            signals.append(("ARI T3→T4", f"{ari_t3_t4:.4f}", "✅ STABLE", "Semestriel"))
        elif ari_t3_t4 > 0.5:
            signals.append(("ARI T3→T4", f"{ari_t3_t4:.4f}", "⚠️ ATTENTION", "Trimestriel"))
        else:
            signals.append(("ARI T3→T4", f"{ari_t3_t4:.4f}", "🚨 DÉRIVE", "Mensuel"))

    if delta_df is not None and len(delta_df) > 0:
        max_delta = delta_df["delta_pct"].abs().max()
        if max_delta > 20:
            signals.append(("Delta CLV max", f"{max_delta:.1f}%", "🚨 DÉRIVE PROFIL", "Immédiat"))
        else:
            signals.append(("Delta CLV max", f"{max_delta:.1f}%", "✅ STABLE", "—"))

    cv_status = "✅ ROBUSTE" if boot_cv < 0.05 else "⚠️ INSTABLE"
    signals.append(("CV Bootstrap", f"{boot_cv:.4f}", cv_status, "—"))


    # ═══════════════════════════════════════════════════════════════════
    #  DOCX GENERATION
    # ═══════════════════════════════════════════════════════════════════
    print(f"\nGénération du DOCX → {OUTPUT_DOCX}")

    doc = docx.Document()
    
    # TITLE
    title = doc.add_heading('RAPPORT DE SIMULATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Analyse de Stabilité Temporelle du Clustering')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # CONFIGURATION
    doc.add_heading("Configuration de la Simulation", level=1)
    doc.add_paragraph(f"Modèle : {best_model_path.name}")
    doc.add_paragraph("Algorithme : UMAP + KMeans")
    doc.add_paragraph(f"k optimal : {BEST_K}")
    doc.add_paragraph(f"Features ({len(FINAL_FEATURES)}) : {', '.join(FINAL_FEATURES)}")
    doc.add_paragraph(f"Clients total : {len(df_labeled):,}")
    
    doc.add_heading("Fenêtres temporelles", level=2)
    for w, cutoff in WINDOWS.items():
        doc.add_paragraph(f"• {w} ≤ {cutoff.date()} : {window_sizes[w]:>7,} clients")
    doc.add_page_break()

    # TABLE: Metrics
    metrics_rows = []
    for w in ["T1", "T2", "T3", "T4"]:
        if w in window_metrics:
            m = window_metrics[w]
            metrics_rows.append([
                w, f"{window_sizes[w]:,}",
                f"{m['silhouette']:.4f}",
                f"{m['davies_bouldin']:.4f}",
                f"{m['calinski_harabasz']:,.0f}",
            ])
    if metrics_rows:
        df_tbl = pd.DataFrame(metrics_rows,
            columns=["Fenêtre", "Clients", "Silhouette", "Davies-Bouldin", "Calinski-Harabasz"])
        add_table(doc, "Métriques de Clustering par Fenêtre", df_tbl)

    # CHART: ARI
    if ari_scores:
        fig, ax = styled_fig()
        labels_ari = list(ari_scores.keys())
        values_ari = list(ari_scores.values())
        colors = [C_GREEN if v > 0.7 else C_ORANGE if v > 0.5 else C_RED for v in values_ari]
        bars = ax.bar(labels_ari, values_ari, color=colors, edgecolor="white", linewidth=1.5, width=0.5)
        ax.axhline(0.7, color=C_GREEN, linestyle="--", linewidth=1.2, label="Seuil stable (0.7)")
        ax.axhline(0.5, color=C_RED, linestyle="--", linewidth=1.2, label="Seuil dérive (0.5)")
        ax.set_ylim(0, 1.05)
        ax.set_ylabel("ARI")
        ax.set_title("Adjusted Rand Index — Fenêtres Consécutives")
        ax.legend()
        for bar, val in zip(bars, values_ari):
            ax.text(bar.get_x() + bar.get_width() / 2, val + 0.02, f"{val:.3f}",
                    ha="center", fontweight="bold")
        plt.tight_layout()
        
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            plt.savefig(tf.name, dpi=150)
            tf.close()
            doc.add_picture(tf.name, width=Inches(6.0))
            os.unlink(tf.name)
        plt.close(fig)

    # CHART: Silhouette
    if silhouette_by_window:
        fig, ax = styled_fig()
        w_list = list(silhouette_by_window.keys())
        s_vals = list(silhouette_by_window.values())
        ax.plot(w_list, s_vals, marker="o", linewidth=2.5, color=C_BLUE, markersize=10)
        ax.fill_between(w_list, s_vals, alpha=0.15, color=C_BLUE)
        for w, v in zip(w_list, s_vals):
            ax.annotate(f"{v:.4f}", (w, v), textcoords="offset points", xytext=(0, 12), ha="center")
        if s_vals:
            drift_thr = max(s_vals) - 0.05
            ax.axhline(drift_thr, color=C_RED, linestyle="--", linewidth=1.2,
                       label=f"Seuil dérive (max − 0.05 = {drift_thr:.3f})")
        ax.set_ylabel("Score de Silhouette")
        ax.set_title("Évolution du Score de Silhouette par Fenêtre")
        ax.legend()
        plt.tight_layout()
        
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            plt.savefig(tf.name, dpi=150)
            tf.close()
            doc.add_picture(tf.name, width=Inches(6.0))
            os.unlink(tf.name)
        plt.close(fig)

    doc.add_page_break()

    # TABLE: Profile
    if "T4" in profile_by_window:
        prof = profile_by_window["T4"].copy()
        display_cols = ["cluster", "n_customers", "pct_customers", "CLV_proxy"] + raw_cols[:3]
        display_cols = [c for c in display_cols if c in prof.columns]
        prof_display = prof[display_cols].copy()
        for c in prof_display.columns:
            if prof_display[c].dtype in [np.float64, np.float32]:
                prof_display[c] = prof_display[c].apply(lambda x: f"{x:,.2f}")
        add_table(doc, "Profil des Clusters — Fenêtre T4 (Dataset complet)", prof_display)

    # TABLE: Delta CLV
    if delta_df is not None and len(delta_df) > 0:
        dd = delta_df.reset_index().copy()
        for c in ["CLV_T3", "CLV_T4"]:
            dd[c] = dd[c].apply(lambda x: f"{x:,.2f}")
        dd["delta_pct"] = delta_df.reset_index()["delta_pct"].apply(lambda x: f"{x:+.2f}%")
        add_table(doc, "Delta CLV Proxy — T3 → T4", dd)

    doc.add_page_break()

    # CHART: Bootstrap
    if bootstrap_sil:
        fig, ax = styled_fig()
        ax.hist(bootstrap_sil, bins=20, color=C_BLUE, edgecolor="white", alpha=0.85)
        ax.axvline(boot_mean, color=C_RED, linewidth=2, label=f"Moyenne = {boot_mean:.4f}")
        ax.axvline(ci95_lo, color=C_ORANGE, linestyle="--", linewidth=1.5,
                   label=f"IC 95% [{ci95_lo:.4f}, {ci95_hi:.4f}]")
        ax.axvline(ci95_hi, color=C_ORANGE, linestyle="--", linewidth=1.5)
        ax.set_xlabel("Score de Silhouette (bootstrap)")
        ax.set_ylabel("Fréquence")
        ax.set_title(f"Distribution Bootstrap (n={N_BOOTSTRAP} iter.) — CV={boot_cv:.4f}")
        ax.legend()
        plt.tight_layout()
        
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            plt.savefig(tf.name, dpi=150)
            tf.close()
            doc.add_picture(tf.name, width=Inches(6.0))
            os.unlink(tf.name)
        plt.close(fig)

    # RECOMMANDATION
    doc.add_heading("Synthèse des signaux de stabilité", level=1)
    for crit, val, status, rec in signals:
        doc.add_paragraph(f"{status} | {crit} = {val} → {rec}")

    freq_map = {"Immédiat": 0, "Mensuel": 1, "Trimestriel": 2, "Semestriel": 3, "—": 4}
    rec_freqs = [r for _, _, _, r in signals if r in freq_map]
    final_rec = min(rec_freqs, key=lambda r: freq_map[r]) if rec_freqs else "Trimestriel"

    doc.add_heading("Recommandation Finale", level=2)
    p = doc.add_paragraph(f"Ré-entraînement recommandé : {final_rec.upper()}")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_heading("Grille de décision", level=3)
    doc.add_paragraph("ARI > 0.7 → ✅ Semestriel")
    doc.add_paragraph("ARI 0.5 – 0.7 → ⚠️ Trimestriel")
    doc.add_paragraph("ARI < 0.5 → 🚨 Mensuel + alerting")
    doc.add_paragraph("Delta CLV > 20% → 🚨 Ré-entraînement immédiat")
    doc.add_paragraph("CV bootstrap < 0.05 → ✅ Modèle robuste")

    doc.save(OUTPUT_DOCX)
    print(f"\n✅ DOCX généré avec succès : {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
